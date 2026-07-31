"""Uploaded CV handling: extract text for the run, archive the original to GCS.

Extraction is synchronous CPU work and GCS uses a sync client; both run via
asyncio.to_thread so the event loop never blocks on a large PDF or upload.
"""
import asyncio
import io
import uuid
from datetime import UTC, datetime

from ada.config import get_settings
from ada.observability import log

ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}
MAX_UPLOAD_BYTES = 5 * 1024 * 1024


class UnsupportedDocument(ValueError):
    """File type or content we can't turn into CV text."""


def _extension(filename: str) -> str:
    dot = filename.rfind(".")
    return filename[dot:].lower() if dot != -1 else ""


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract(filename: str, data: bytes) -> str:
    ext = _extension(filename)
    if ext == ".pdf":
        text = _extract_pdf(data)
    elif ext == ".docx":
        text = _extract_docx(data)
    elif ext == ".txt":
        text = data.decode("utf-8", errors="replace")
    else:
        raise UnsupportedDocument(f"unsupported file type: {ext or 'no extension'}")
    text = text.strip()
    if len(text) < 30:
        raise UnsupportedDocument(
            "couldn't read enough text from that file — is it a scan or an image?"
        )
    return text


def _store(user_id: str, filename: str, data: bytes, content_type: str | None) -> str | None:
    bucket_name = get_settings().gcs_bucket
    if not bucket_name:
        return None
    from google.cloud import storage

    stamp = datetime.now(UTC).strftime("%Y%m%d")
    blob_name = f"cv/{user_id}/{stamp}-{uuid.uuid4().hex[:8]}-{filename[-80:]}"
    client = storage.Client()
    blob = client.bucket(bucket_name).blob(blob_name)
    blob.upload_from_string(data, content_type=content_type or "application/octet-stream")
    return f"gs://{bucket_name}/{blob_name}"


async def extract_cv_text(filename: str, data: bytes) -> str:
    """Extract text from an uploaded CV without archiving it — for the public,
    unauthenticated assessment, where there's no user to own a stored document.
    Raises UnsupportedDocument on bad input."""
    return await asyncio.to_thread(_extract, filename, data)


async def process_cv_upload(
    user_id: str, filename: str, data: bytes, content_type: str | None
) -> tuple[str, str | None]:
    """Returns (extracted_text, gcs_uri). Raises UnsupportedDocument on bad input."""
    text = await asyncio.to_thread(_extract, filename, data)
    try:
        gcs_uri = await asyncio.to_thread(_store, user_id, filename, data, content_type)
    except Exception as exc:  # archive failure shouldn't lose the extracted text
        log.warning("cv_archive_failed", error=str(exc), user_id=user_id)
        gcs_uri = None
    if gcs_uri:
        log.info("cv_archived", user_id=user_id, gcs_uri=gcs_uri, bytes=len(data))
    return text, gcs_uri
