"""One-click apply: build applicant answers, render the CV, drive the ATS submit."""
import asyncio
import io
import re

from ada.config import get_settings
from ada.db.models import ApplicationStatus, Job, Profile, RunStatus, User
from ada.db.repositories import (
    ApplicationRepository,
    JobRepository,
    ProfileRepository,
    RunRepository,
)
from ada.db.session import _session_factory
from ada.observability import log
from ada.services.ats import ApplicantAnswers, FormPlan, split_name
from ada.services.ats import ashby as ats_ashby
from ada.services.ats import greenhouse as ats_greenhouse
from ada.services.ats import lever as ats_lever

SUPPORTED_SOURCES = {"greenhouse", "lever", "ashby"}

_browser_slots: asyncio.Semaphore | None = None


def _browser_gate() -> asyncio.Semaphore:
    global _browser_slots
    if _browser_slots is None:
        _browser_slots = asyncio.Semaphore(get_settings().apply_max_concurrency)
    return _browser_slots


class ApplyPrecondition(ValueError):
    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code


def safe_filename(name: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "-", name).strip("-.")
    return cleaned[:80] or "applicant"


def cv_docx_bytes(cv_markdown: str) -> bytes:
    from docx import Document

    doc = Document()
    for raw in cv_markdown.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        stripped = line.lstrip("#").strip()
        if line.startswith("# "):
            doc.add_heading(stripped, level=1)
        elif line.startswith("## "):
            doc.add_heading(stripped, level=2)
        elif line.startswith("### "):
            doc.add_heading(stripped, level=3)
        elif line.lstrip().startswith(("- ", "* ")):
            doc.add_paragraph(line.lstrip()[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line.replace("**", ""))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_answers(user: User, profile: Profile | None, cv_markdown: str) -> ApplicantAnswers:
    if profile is None or not (profile.full_name or "").strip():
        raise ApplyPrecondition(
            "identity_required", "Add your full name (and phone) so Ada can apply for you."
        )
    full_name = (profile.full_name or "").strip()
    first, last = split_name(full_name)
    return ApplicantAnswers(
        full_name=full_name,
        first_name=first,
        last_name=last,
        email=user.email,
        phone=(profile.phone or "").strip() or None,
        linkedin_url=profile.linkedin_url,
        cv_filename=f"{safe_filename(full_name)}-CV.docx",
        cv_bytes=cv_docx_bytes(cv_markdown),
    )


def build_plan(job: Job, answers: ApplicantAnswers) -> FormPlan:
    if not job.url:
        raise ApplyPrecondition("no_url", "This listing has no application page on file.")
    if job.source == ats_greenhouse.SOURCE:
        return ats_greenhouse.build_form_plan(job.url, answers)
    if job.source == ats_lever.SOURCE:
        return ats_lever.build_form_plan(job.url, answers)
    if job.source == ats_ashby.SOURCE:
        return ats_ashby.build_form_plan(job.url, answers)
    raise ApplyPrecondition("unsupported_source", "No deterministic plan for this source.")


async def latest_cv_markdown(runs: RunRepository, user_id: str) -> tuple[str, str] | None:
    for run in await runs.list_by_user(user_id):
        if run.status == RunStatus.COMPLETE and run.rewritten_cv:
            return run.rewritten_cv, run.id
    return None


async def run_submission(application_id: str) -> None:
    from ada.services.ats.executor import execute_plan
    from ada.services.ats.generic import submit_generic

    async with _session_factory() as session:
        applications = ApplicationRepository(session)
        application = await applications.get(application_id)
        if application is None:
            return
        job = await JobRepository(session).get(application.job_id)
        profile = await ProfileRepository(session).get(application.user_id)
        user = await session.get(User, application.user_id)
        run = await RunRepository(session).get(application.run_id) if application.run_id else None
        try:
            if user is None or job is None or run is None or not run.rewritten_cv:
                raise ApplyPrecondition("missing_data", "Application inputs went missing.")
            if not job.url:
                raise ApplyPrecondition("no_url", "This listing has no application page on file.")
            answers = build_answers(user, profile, run.rewritten_cv)
            plan = build_plan(job, answers) if job.source in SUPPORTED_SOURCES else None
            job_url = job.url
        except ApplyPrecondition as exc:
            await applications.set_status(
                application_id, ApplicationStatus.NEEDS_ATTENTION, detail=str(exc)
            )
            return
    try:
        async with _browser_gate():
            if plan is not None:
                outcome = await execute_plan(plan, answers)
            else:
                outcome = await submit_generic(job_url, answers)
    except Exception as exc:  # noqa: BLE001 — any submit failure must land in status, not logs alone
        log.error("apply_submit_crashed", application_id=application_id, error=str(exc))
        outcome = None
    async with _session_factory() as session:
        applications = ApplicationRepository(session)
        application = await applications.get(application_id)
        job = await JobRepository(session).get(application.job_id) if application else None
        role = job.title if job else "the role"
        if outcome is None:
            await applications.set_status(
                application_id, ApplicationStatus.FAILED,
                detail="Something broke on our side — we'll look into it. Nothing was submitted.",
            )
            _note = ("run_failed", "Application couldn't be sent",
                     f"We hit a snag applying to {role} — nothing was submitted. Try again.")
        elif outcome.status == "submitted":
            await applications.set_status(application_id, ApplicationStatus.SUBMITTED)
            log.info("apply_submitted", application_id=application_id)
            _note = ("application_submitted", "Application submitted",
                     f"Ada submitted your application to {role}.")
        else:
            await applications.set_status(
                application_id, ApplicationStatus.NEEDS_ATTENTION, detail=outcome.detail
            )
            _note = ("application_attention", "An application needs you",
                     f"Your application to {role} needs a step only you can do: "
                     f"{outcome.detail}")
        if application is not None:
            from ada.services.notify import notify

            await notify(
                application.user_id, kind=_note[0], title=_note[1], body=_note[2],
                link="/app/applications",
            )


def is_supported(job: Job) -> bool:
    return bool(job.url)


async def recover_stuck_applications() -> int:
    """Applications left PREPARING past the window had their in-process submit lost
    (instance terminated). Flip them to NEEDS_ATTENTION so the user can retry."""
    threshold = get_settings().apply_stuck_seconds
    async with _session_factory() as session:
        applications = ApplicationRepository(session)
        stuck = await applications.find_stuck(threshold)
        for application_id in stuck:
            await applications.set_status(
                application_id,
                ApplicationStatus.NEEDS_ATTENTION,
                detail="This application was interrupted before finishing — tap Apply to retry.",
            )
    return len(stuck)
